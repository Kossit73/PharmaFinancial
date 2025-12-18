"""Shared report generation utilities for financial models."""
from __future__ import annotations

import json
import math
import re
from dataclasses import dataclass
from datetime import UTC, datetime
from io import BytesIO, StringIO
from typing import Any, Iterable, List, Mapping, MutableMapping, Sequence, Tuple
from xml.sax.saxutils import escape as xml_escape
from zipfile import ZipFile

try:  # pragma: no cover - optional dependency
    import pandas as pd  # type: ignore
except Exception:  # pragma: no cover - pandas optional in minimal environments
    pd = None  # type: ignore

try:  # pragma: no cover - optional dependency
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - python-docx may not be installed
    Document = None  # type: ignore

try:  # pragma: no cover - optional dependency
    from fpdf import FPDF  # type: ignore
except Exception:  # pragma: no cover - FPDF may not be installed
    FPDF = None  # type: ignore

from .model import FinancialModel, FinancialOutputs
from .table import Table, build_table

JSON_MIME = "application/json"
CSV_MIME = "text/csv"
EXCEL_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
WORD_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"

REPORT_FORMATS = ["PDF", "Word", "Excel", "CSV", "JSON"]


class ReportGenerationError(RuntimeError):
    """Raised when the requested report cannot be produced."""


@dataclass
class ReportTable:
    """Container describing an individual table within a report."""

    title: str
    data: Any
    note: str | None = None


@dataclass
class ReportSection:
    """Represents a logical report section grouping multiple tables."""

    title: str
    tables: List[ReportTable]
    notes: List[str] | None = None


def _table_from_df(df: Any, *, index_name: str = "Year") -> Table | Any:
    """Convert a pandas DataFrame into a Table when possible, otherwise return original."""

    if df is None:
        return []
    if isinstance(df, Table):
        return df
    if pd is not None and hasattr(df, "index") and hasattr(df, "columns"):
        return build_table(list(df.index), {c: df[c].tolist() for c in df.columns}, index_name=index_name)
    return df


def collect_report_sections(model: FinancialModel, outputs: FinancialOutputs) -> List[ReportSection]:
    """Assemble pharma report sections spanning all dashboards in the required order."""

    sections: List[ReportSection] = []

    key_tables: List[ReportTable] = [
        ReportTable("Summary Metrics", outputs.summary_metrics),
        ReportTable("Goal Seek", outputs.goal_seek),
    ]

    try:
        key_tables.append(ReportTable("Working Capital Schedule", model.working_capital_schedule()))
    except Exception as exc:  # pragma: no cover - defensive guard for runtime calculations
        key_tables.append(ReportTable("Working Capital Schedule", [], note=f"Unavailable: {exc}"))

    try:
        key_tables.append(ReportTable("Inventory Schedule", model.inventory_schedule()))
    except Exception as exc:  # pragma: no cover - defensive guard for runtime calculations
        key_tables.append(ReportTable("Inventory Schedule", [], note=f"Unavailable: {exc}"))

    ai_notes: List[str] = []
    if outputs.ai_insights is not None:
        insight = outputs.ai_insights
        if insight.ml_forecast is not None:
            key_tables.append(ReportTable("AI Revenue Forecast", insight.ml_forecast))
        if insight.generative_summary:
            ai_notes.append(insight.generative_summary)
        if insight.metadata:
            provider = insight.metadata.get("provider")
            model_name = insight.metadata.get("model")
            status = insight.metadata.get("status")
            details = ", ".join(str(value) for value in [provider, model_name, status] if value not in (None, ""))
            if details:
                ai_notes.append(f"AI configuration: {details}")

    sections.append(ReportSection("Key Metrics Dashboard", key_tables, notes=ai_notes or None))

    perf_tables: List[ReportTable] = [
        ReportTable(
            "Statement of Financial Performance",
            outputs.income_statement.rounded(0, exclude_keywords=("Margin", "Return")),
        ),
    ]
    try:
        perf_tables.append(ReportTable("Gross Revenue Schedule", model.revenue_schedule()))
    except Exception as exc:  # pragma: no cover - defensive guard
        perf_tables.append(ReportTable("Gross Revenue Schedule", [], note=f"Unavailable: {exc}"))

    try:
        perf_tables.append(ReportTable("Total Expenses Schedule", model.cost_structure()))
    except Exception as exc:  # pragma: no cover - defensive guard
        perf_tables.append(ReportTable("Total Expenses Schedule", [], note=f"Unavailable: {exc}"))

    sections.append(ReportSection("Financial Performance", perf_tables))

    break_even_tables = [
        ReportTable("Break-even Analysis", outputs.break_even),
        ReportTable("Payback Schedule", outputs.payback),
        ReportTable("Discounted Payback Schedule", outputs.discounted_payback),
    ]
    sections.append(ReportSection("Break-even & Payback", break_even_tables))

    sections.append(
        ReportSection(
            "Financial Position",
            [
                ReportTable(
                    "Statement of Financial Position",
                    outputs.balance_sheet.rounded(0),
                )
            ],
        )
    )

    sections.append(
        ReportSection(
            "Cash Flow Statement",
            [
                ReportTable(
                    "Statement of Cash Flows",
                    outputs.cash_flow.rounded(0),
                )
            ],
        )
    )

    sensitivity_tables: List[ReportTable] = []
    for name, table in outputs.sensitivity_results.items():
        label = name.replace("_", " ").title()
        sensitivity_tables.append(ReportTable(f"Sensitivity: {label}", table))
    if not sensitivity_tables:
        sensitivity_tables.append(ReportTable("Sensitivity Analysis", [], note="No sensitivity configurations provided."))
    sections.append(ReportSection("Sensitivity Analysis", sensitivity_tables))

    scenario_tables: List[ReportTable] = []
    for name, table in outputs.scenario_results.items():
        scenario_tables.append(ReportTable(f"Scenario: {name}", table))
    for key, result in outputs.scenario_tool_results.items():
        label = key.replace("_", " ").title()
        scenario_tables.append(ReportTable(f"Scenario Tool: {label}", result.rows, note=result.interpretation))
    if not scenario_tables:
        scenario_tables.append(ReportTable("Scenario / IFs Analysis", [], note="No scenario configurations provided."))
    sections.append(ReportSection("Scenario / IFs Analysis", scenario_tables))

    sections.append(ReportSection("Monte Carlo Simulation", [ReportTable("Monte Carlo Simulation", outputs.monte_carlo)]))

    return sections


def collect_biotech_report_sections(result: Any) -> List[ReportSection]:
    """Build report sections for biotech valuation results."""

    sections: List[ReportSection] = []
    dcf_table = _table_from_df(getattr(result, "dcf_table", None))
    consolidated = _table_from_df(getattr(result, "consolidated", None))
    overview_tables = []
    if getattr(result, "rnpv", None) is not None:
        overview_tables.append(ReportTable("rNPV", [{"rNPV": result.rnpv}]))
    if dcf_table:
        overview_tables.append(ReportTable("Discounted Cash Flows", dcf_table))
    if consolidated:
        overview_tables.append(ReportTable("Consolidated Cash Flows", consolidated))
    sections.append(ReportSection("Biotech Valuation Overview", overview_tables))

    per_product_tables: List[ReportTable] = []
    for name, df in getattr(result, "per_product", {}).items():
        per_product_tables.append(ReportTable(f"Per-product Cashflows: {name}", _table_from_df(df)))
    if per_product_tables:
        sections.append(ReportSection("Per-product Cashflows", per_product_tables))

    per_product_prob_tables: List[ReportTable] = []
    for name, df in getattr(result, "per_product_prob", {}).items():
        per_product_prob_tables.append(ReportTable(f"Probability-weighted Cashflows: {name}", _table_from_df(df)))
    if per_product_prob_tables:
        sections.append(ReportSection("Probability-weighted Cashflows", per_product_prob_tables))

    return sections


def generate_report(
    sections: Sequence[ReportSection],
    format_name: str,
    *,
    report_name: str = "financial_report",
    report_title: str | None = None,
) -> Tuple[bytes, str, str]:
    """Generate a consolidated report in the requested format."""

    if not sections:
        raise ReportGenerationError("No sections available to generate a report.")

    fmt = format_name.strip().lower()
    if fmt not in {f.lower() for f in REPORT_FORMATS}:
        raise ReportGenerationError(f"Unsupported report format: {format_name}")

    timestamp = datetime.now(UTC).strftime("%Y%m%d%H%M%S")

    if fmt == "json":
        data = _build_json(sections)
        filename = f"{report_name}_{timestamp}.json"
        return data, JSON_MIME, filename

    if fmt == "csv":
        data = _build_csv(sections)
        filename = f"{report_name}_{timestamp}.csv"
        return data, CSV_MIME, filename

    if fmt == "excel":
        data = _build_excel(sections)
        filename = f"{report_name}_{timestamp}.xlsx"
        return data, EXCEL_MIME, filename

    if fmt == "word":
        data = _build_word(sections, title=report_title)
        filename = f"{report_name}_{timestamp}.docx"
        return data, WORD_MIME, filename

    if fmt == "pdf":
        data = _build_pdf(sections, title=report_title)
        filename = f"{report_name}_{timestamp}.pdf"
        return data, PDF_MIME, filename

    raise ReportGenerationError(f"Unsupported report format: {format_name}")


# ---------------------------------------------------------------------------
# format builders
# ---------------------------------------------------------------------------

def _build_json(sections: Sequence[ReportSection]) -> bytes:
    payload: List[MutableMapping[str, Any]] = []
    for section in sections:
        entry: MutableMapping[str, Any] = {
            "title": section.title,
            "tables": [],
        }
        if section.notes:
            entry["notes"] = list(section.notes)
        for table in section.tables:
            entry["tables"].append(
                {
                    "title": table.title,
                    "rows": _table_to_rows(table.data),
                    **({"note": table.note} if table.note else {}),
                }
            )
        payload.append(entry)
    return json.dumps({"generated": datetime.now(UTC).isoformat(), "sections": payload}, indent=2).encode("utf-8")


def _build_csv(sections: Sequence[ReportSection]) -> bytes:
    buffer = StringIO()
    for section in sections:
        buffer.write(f"# Section: {section.title}\n")
        if section.notes:
            for note in section.notes:
                buffer.write(f"# Note: {note}\n")
        for table in section.tables:
            buffer.write(f"## Table: {table.title}\n")
            if table.note:
                buffer.write(f"## Note: {table.note}\n")
            rows = _table_to_rows(table.data)
            if not rows:
                buffer.write("(no data)\n\n")
                continue
            columns = list(rows[0].keys())
            buffer.write(",".join(_escape_csv(value) for value in columns) + "\n")
            for row in rows:
                buffer.write(",".join(_escape_csv(row.get(column, "")) for column in columns) + "\n")
            buffer.write("\n")
        buffer.write("\n")
    return buffer.getvalue().encode("utf-8")


def _build_excel(sections: Sequence[ReportSection]) -> bytes:
    entries = _collect_sheet_entries(sections)

    if pd is not None:
        buffer = BytesIO()
        try:
            with pd.ExcelWriter(buffer) as writer:  # type: ignore[arg-type]
                for entry in entries:
                    frame = _rows_to_dataframe(entry["rows"])
                    if frame is None:
                        raise ValueError("pandas unavailable")
                    frame = frame.copy()
                    note = entry.get("note")
                    if note:
                        note_col = "Notes"
                        if frame.empty:
                            frame = pd.DataFrame({note_col: [note]})
                        else:
                            if note_col not in frame.columns:
                                frame[note_col] = ""
                            frame.iloc[0, frame.columns.get_loc(note_col)] = note
                    frame.to_excel(writer, sheet_name=entry["sheet_name"], index=False)
            return buffer.getvalue()
        except Exception:  # pragma: no cover - fall back to pure-Python writer
            pass

    return _build_excel_fallback(entries)


def _build_word(sections: Sequence[ReportSection], *, title: str | None = None) -> bytes:
    blocks = _collect_report_blocks(sections, title=title)

    if Document is not None:
        document = Document()
        for kind, text in blocks:
            if kind == "title":
                document.add_heading(text, level=1)
            elif kind == "subtitle":
                document.add_paragraph(text, style="Subtitle")
            elif kind == "heading1":
                document.add_heading(text, level=2)
            elif kind == "heading2":
                document.add_heading(text, level=3)
            elif kind == "note":
                document.add_paragraph(text, style="Intense Quote")
            else:
                document.add_paragraph(text)
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    return _build_word_fallback(blocks)


def _build_pdf(sections: Sequence[ReportSection], *, title: str | None = None) -> bytes:
    blocks = _collect_report_blocks(sections, title=title)

    if FPDF is not None:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        for kind, text in blocks:
            if kind == "title":
                pdf.set_font("Arial", "B", 16)
                pdf.cell(0, 10, text, ln=True)
            elif kind == "subtitle":
                pdf.set_font("Arial", "", 10)
                pdf.cell(0, 8, text, ln=True)
            elif kind == "heading1":
                pdf.ln(4)
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 8, text, ln=True)
            elif kind == "heading2":
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 7, text, ln=True)
            elif kind == "note":
                _pdf_multiline(pdf, text)
            else:
                _pdf_multiline(pdf, text)
        return bytes(pdf.output(dest="S").encode("latin-1"))

    return _build_pdf_fallback(blocks)


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------


def _collect_report_blocks(sections: Sequence[ReportSection], *, title: str | None = None) -> List[Tuple[str, str]]:
    """Flatten sections into typed text blocks for downstream exporters."""

    blocks: List[Tuple[str, str]] = [
        ("title", title or "Financial Model Report"),
        ("subtitle", f"Generated on {datetime.now(UTC).isoformat()} UTC"),
    ]

    for section in sections:
        blocks.append(("heading1", section.title))
        if section.notes:
            for note in section.notes:
                blocks.append(("note", f"Note: {note}"))
        for table in section.tables:
            blocks.append(("heading2", table.title))
            if table.note:
                blocks.append(("note", f"Note: {table.note}"))
            rows = _table_to_rows(table.data)
            if not rows:
                blocks.append(("body", "No data available."))
                continue
            columns = _ordered_columns(rows)
            if columns:
                blocks.append(("body", " | ".join(columns)))
                for row in rows:
                    values = [_stringify(row.get(column, "")) for column in columns]
                    blocks.append(("body", " | ".join(values)))
            else:
                for row in rows:
                    values = [_stringify(value) for value in row.values()]
                    blocks.append(("body", " | ".join(values) or "No data available."))

    return blocks


def _collect_sheet_entries(sections: Sequence[ReportSection]) -> List[dict[str, Any]]:
    """Prepare sheet metadata shared by the Excel exporters."""

    tracker: dict[str, int] = {}
    entries: List[dict[str, Any]] = []

    for section in sections:
        for table in section.tables:
            rows = _table_to_rows(table.data)
            sheet_name = _sheet_name(section.title, table.title, tracker)
            entries.append({"sheet_name": sheet_name, "rows": rows, "note": table.note})
        if section.notes:
            sheet_name = _sheet_name(section.title, "Notes", tracker)
            note_rows = [{"Notes": note} for note in section.notes]
            entries.append({"sheet_name": sheet_name, "rows": note_rows, "note": None})

    if not entries:
        entries.append(
            {
                "sheet_name": "Report",
                "rows": [{"Message": "No data available."}],
                "note": None,
            }
        )

    return entries


def _rows_to_dataframe(rows: List[Mapping[str, Any]]):
    """Convert a list of dictionaries to a pandas DataFrame when available."""

    if pd is None:
        return None
    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


def _build_excel_fallback(entries: Sequence[Mapping[str, Any]]) -> bytes:
    """Create a minimal XLSX archive without third-party dependencies."""

    sheets: List[Tuple[str, List[List[str]]]] = []
    for entry in entries:
        rows = entry.get("rows", [])
        note = entry.get("note")
        matrix = _rows_to_matrix(rows, note)
        sheets.append((entry["sheet_name"], matrix))

    return _write_xlsx_archive(sheets)


def _rows_to_matrix(rows: Sequence[Mapping[str, Any]], note: Any) -> List[List[str]]:
    """Normalise rows into a rectangular matrix of strings."""

    columns = _ordered_columns(rows)
    matrix: List[List[str]] = []

    if columns:
        matrix.append(columns)
        for row in rows:
            matrix.append([_stringify(row.get(column, "")) for column in columns])
    else:
        matrix.append(["Value"])
        if rows:
            for row in rows:
                value = next(iter(row.values()), "")
                matrix.append([_stringify(value)])
        else:
            matrix.append(["No data available."])
            note = None

    if note:
        padding = [""] * (len(matrix[0]) - 1)
        matrix.append([f"Note: {note}", *padding])

    # Ensure rectangular shape
    width = max(len(row) for row in matrix)
    for row in matrix:
        if len(row) < width:
            row.extend([""] * (width - len(row)))

    return matrix


def _write_xlsx_archive(sheets: Sequence[Tuple[str, List[List[str]]]]) -> bytes:
    buffer = BytesIO()
    sheet_count = max(len(sheets), 1)
    sheet_names = [name for name, _ in sheets] or ["Report"]

    with ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", _excel_content_types(sheet_count))
        archive.writestr("_rels/.rels", _excel_root_rels())
        archive.writestr("docProps/core.xml", _excel_core_props())
        archive.writestr("docProps/app.xml", _excel_app_props(sheet_names))
        archive.writestr("xl/_rels/workbook.xml.rels", _excel_workbook_rels(sheet_count))
        archive.writestr("xl/workbook.xml", _excel_workbook_xml(sheet_names))
        archive.writestr("xl/styles.xml", _excel_styles_xml())

        if sheets:
            for index, (_, rows) in enumerate(sheets, start=1):
                archive.writestr(f"xl/worksheets/sheet{index}.xml", _excel_sheet_xml(rows))
        else:
            archive.writestr("xl/worksheets/sheet1.xml", _excel_sheet_xml([["Message"], ["No data available."]]))

    return buffer.getvalue()


def _excel_sheet_xml(rows: Sequence[Sequence[str]]) -> str:
    max_cols = max(len(row) for row in rows) if rows else 1
    parts = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">',
        "<sheetData>",
    ]

    for r_idx, row in enumerate(rows, start=1):
        parts.append(f'<row r="{r_idx}">')
        padded = list(row) + [""] * (max_cols - len(row))
        for c_idx, value in enumerate(padded, start=1):
            column = _excel_column_letter(c_idx)
            escaped = xml_escape(_stringify(value))
            parts.append(
                f'<c r="{column}{r_idx}" t="inlineStr"><is><t xml:space="preserve">{escaped}</t></is></c>'
            )
        parts.append("</row>")

    parts.append("</sheetData></worksheet>")
    return "\n".join(parts)


def _excel_content_types(n_sheets: int) -> str:
    overrides = "".join(
        f'<Override PartName="/xl/worksheets/sheet{i}.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        for i in range(1, n_sheets + 1)
    )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        f"{overrides}"
        '<Override PartName="/xl/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>'
        '<Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>'
        '<Override PartName="/xl/sharedStrings.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sharedStrings+xml"/>'
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
        '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
        "</Types>"
    )


def _excel_root_rels() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
        '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>'
        '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>'
        "</Relationships>"
    )


def _excel_workbook_rels(n_sheets: int) -> str:
    entries = "".join(
        f'<Relationship Id="rId{i}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet{i}.xml"/>'
        for i in range(1, n_sheets + 1)
    )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        f"{entries}"
        '<Relationship Id="rId{last}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
        '<Relationship Id="rId{ss}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/sharedStrings" Target="sharedStrings.xml"/>'
    ).format(last=n_sheets + 1, ss=n_sheets + 2)


def _excel_workbook_xml(sheet_names: Sequence[str]) -> str:
    sheets_xml = "".join(
        f'<sheet name="{xml_escape(name)}" sheetId="{idx}" r:id="rId{idx}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
        for idx, name in enumerate(sheet_names, start=1)
    )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        "<sheets>"
        f"{sheets_xml}"
        "</sheets>"
        "</workbook>"
    )


def _excel_styles_xml() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        "<fonts count=\"1\"><font><sz val=\"11\"/><color theme=\"1\"/><name val=\"Calibri\"/></font></fonts>"
        "<fills count=\"1\"><fill><patternFill patternType=\"none\"/></fill></fills>"
        "<borders count=\"1\"><border><left/><right/><top/><bottom/><diagonal/></border></borders>"
        "<cellStyleXfs count=\"1\"><xf numFmtId=\"0\" fontId=\"0\" fillId=\"0\" borderId=\"0\"/></cellStyleXfs>"
        "<cellXfs count=\"1\"><xf numFmtId=\"0\" fontId=\"0\" fillId=\"0\" borderId=\"0\" xfId=\"0\"/></cellXfs>"
        "</styleSheet>"
    )


def _excel_app_props(sheet_names: Sequence[str]) -> str:
    names = "".join(f"<vt:lpstr>{xml_escape(name)}</vt:lpstr>" for name in sheet_names)
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        "<Properties xmlns=\"http://schemas.openxmlformats.org/officeDocument/2006/extended-properties\" "
        "xmlns:vt=\"http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes\">"
        "<Application>Python</Application>"
        "<DocSecurity>0</DocSecurity>"
        "<ScaleCrop>false</ScaleCrop>"
        "<HeadingPairs><vt:vector size=\"2\" baseType=\"variant\"><vt:variant><vt:lpstr>Worksheets</vt:lpstr></vt:variant><vt:variant><vt:i4>"
        f"{len(sheet_names)}</vt:i4></vt:variant></vt:vector></HeadingPairs>"
        "<TitlesOfParts><vt:vector size=\"{n}\" baseType=\"lpstr\">{names}</vt:vector></TitlesOfParts>"
        "<Manager/>"
        "<Company/>"
        "<LinksUpToDate>false</LinksUpToDate>"
        "<SharedDoc>false</SharedDoc>"
        "<HyperlinksChanged>false</HyperlinksChanged>"
        "<AppVersion>16.0300</AppVersion>"
        "</Properties>"
    ).format(n=len(sheet_names), names=names)


def _excel_core_props() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        "<cp:coreProperties xmlns:cp=\"http://schemas.openxmlformats.org/package/2006/metadata/core-properties\" "
        "xmlns:dc=\"http://purl.org/dc/elements/1.1/\" xmlns:dcterms=\"http://purl.org/dc/terms/\" "
        "xmlns:dcmitype=\"http://purl.org/dc/dcmitype/\" xmlns:xsi=\"http://www.w3.org/2001/XMLSchema-instance\">"
        f"<dcterms:created xsi:type=\"dcterms:W3CDTF\">{datetime.now(UTC).isoformat()}</dcterms:created>"
        "<dc:creator>financial_models</dc:creator>"
        "<cp:revision>1</cp:revision>"
        "</cp:coreProperties>"
    )


def _excel_column_letter(n: int) -> str:
    string = ""
    while n > 0:
        n, remainder = divmod(n - 1, 26)
        string = chr(65 + remainder) + string
    return string


def _excel_sheet_xml_value(value: str) -> str:
    escaped = xml_escape(value)
    return f"<t xml:space=\"preserve\">{escaped}</t>"


def _escape_csv(value: Any) -> str:
    text = _stringify(value)
    if any(ch in text for ch in [",", "\n", '"']):
        return '"' + text.replace('"', '""') + '"'
    return text


def _stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        if math.isnan(value):
            return ""
        if math.isinf(value):
            return "∞" if value > 0 else "-∞"
    return str(value)


def _ordered_columns(rows: Sequence[Mapping[str, Any]]) -> List[str]:
    if not rows:
        return []
    columns = list(rows[0].keys())
    return columns


def _sheet_name(section_title: str, table_title: str, tracker: MutableMapping[str, int]) -> str:
    base = re.sub(r"[^A-Za-z0-9]", "_", f"{section_title}_{table_title}").strip("_") or "Sheet"
    max_len = 31
    count = tracker.get(base, 0)
    tracker[base] = count + 1
    suffix = f"_{count}" if count else ""
    allowed = max_len - len(suffix)
    base = base[:allowed] if allowed > 0 else base[:max_len]
    name = f"{base}{suffix}"
    return name[:max_len]


def _table_to_rows(data: Any) -> List[MutableMapping[str, Any]]:
    """Normalise different table structures into a list of row dictionaries."""

    if data is None:
        return []

    if isinstance(data, Table):
        rows: List[MutableMapping[str, Any]] = []
        columns = data.columns()
        for idx, _ in enumerate(data.index):
            row = {data.index_name: data.index[idx]}
            for column in columns:
                row[column] = data.data[column][idx]
            rows.append(row)
        return rows

    if pd is not None and isinstance(data, pd.DataFrame):
        rows = data.reset_index().to_dict(orient="records")  # type: ignore[attr-defined]
        return [dict(row) for row in rows]

    if isinstance(data, Mapping):
        return [dict(data)]

    if isinstance(data, (list, tuple)):
        if not data:
            return []
        if all(isinstance(item, Mapping) for item in data):
            return [dict(item) for item in data]  # type: ignore[arg-type]
        # Coerce positional rows into a single-column table
        return [{"value": item} for item in data]  # type: ignore[list-item]

    return [{"value": data}]


def _pdf_multiline(pdf: Any, text: str) -> None:
    pdf.set_font("Arial", size=10)
    for line in text.splitlines() or [""]:
        pdf.multi_cell(0, 5, line)


def _build_word_fallback(blocks: Sequence[Tuple[str, str]]) -> bytes:
    buffer = StringIO()
    for kind, text in blocks:
        prefix = {
            "title": "# ",
            "subtitle": "## ",
            "heading1": "### ",
            "heading2": "#### ",
            "note": "> ",
        }.get(kind, "")
        buffer.write(f"{prefix}{text}\n\n")
    return buffer.getvalue().encode("utf-8")


def _build_pdf_fallback(blocks: Sequence[Tuple[str, str]]) -> bytes:
    content = "\n".join(text for _, text in blocks)
    return content.encode("utf-8")


__all__ = [
    "ReportGenerationError",
    "ReportSection",
    "ReportTable",
    "collect_report_sections",
    "collect_biotech_report_sections",
    "generate_report",
    "REPORT_FORMATS",
    "JSON_MIME",
    "CSV_MIME",
    "EXCEL_MIME",
    "WORD_MIME",
    "PDF_MIME",
]
