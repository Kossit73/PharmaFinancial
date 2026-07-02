import importlib
import json
import sys
import types
import unittest
from io import BytesIO
from pathlib import Path

try:  # pragma: no cover - optional dependency check
    from docx import Document
except Exception:  # pragma: no cover - import guard for missing optional dependency
    Document = None  # type: ignore

try:  # pragma: no cover - optional dependency check
    from fpdf import FPDF
except Exception:  # pragma: no cover - import guard for missing optional dependency
    FPDF = None  # type: ignore

try:  # pragma: no cover - optional dependency check
    from openpyxl import Workbook
except Exception:  # pragma: no cover - import guard for missing optional dependency
    Workbook = None  # type: ignore


class _NoOp:
    def __call__(self, *args, **kwargs):  # pragma: no cover - defensive noop
        return None

    def __getattr__(self, _name):  # pragma: no cover - defensive noop
        return self


class DummyStreamlit(types.ModuleType):
    def __init__(self):
        super().__init__("streamlit")
        self._no_op = _NoOp()
        self.calls: list[str] = []
        self.session_state: dict = {}

    def rerun(self):  # pragma: no cover - behaviour exercised in tests
        self.calls.append("rerun")
        raise RuntimeError("rerun not available")

    def experimental_rerun(self):  # pragma: no cover - exercised in tests
        self.calls.append("experimental_rerun")
        raise RuntimeError("experimental rerun not available")

    def __getattr__(self, name):  # pragma: no cover - fallback for unused attrs
        return getattr(self._no_op, name, self._no_op)


class RerunHelperTest(unittest.TestCase):
    def setUp(self):
        self.original_streamlit = sys.modules.get("streamlit")
        self.original_runtime = sys.modules.get("streamlit.runtime")

        stub = DummyStreamlit()
        runtime = types.ModuleType("streamlit.runtime")
        runtime.exists = lambda: False  # type: ignore[attr-defined]

        sys.modules["streamlit"] = stub
        sys.modules["streamlit.runtime"] = runtime

        if "pharma_financial.app" in sys.modules:
            del sys.modules["pharma_financial.app"]

        importlib.invalidate_caches()
        self.stub = stub
        self.app = importlib.import_module("pharma_financial.app")
        self.app._INPUT_CACHE.clear()
        self.app._MODEL_CACHE.clear()

    def tearDown(self):
        if self.original_streamlit is None:
            sys.modules.pop("streamlit", None)
        else:
            sys.modules["streamlit"] = self.original_streamlit

        if self.original_runtime is None:
            sys.modules.pop("streamlit.runtime", None)
        else:
            sys.modules["streamlit.runtime"] = self.original_runtime

        if "pharma_financial.app" in sys.modules:
            del sys.modules["pharma_financial.app"]

    def test_rerun_helper_handles_missing_runtime(self):
        self.app._rerun()
        self.assertEqual(self.stub.calls, ["rerun", "experimental_rerun"])

    def test_summary_metric_returns_none_for_non_finite_values(self):
        outputs = types.SimpleNamespace(
            summary_metrics=self.app.Table(
                ["IRR", "Payback Period"],
                {"Value": [float("nan"), float("inf")]},
                index_name="Metric",
            )
        )

        self.assertIsNone(self.app._summary_metric(outputs, "IRR"))
        self.assertIsNone(self.app._summary_metric(outputs, "Payback Period"))

    def test_summary_metric_returns_default_case_values(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(
                encoding="utf-8"
            )
        )
        inputs = self.app.parse_inputs(payload)
        model = self.app.FinancialModel(inputs)
        outputs = model.run_core()

        irr = self.app._summary_metric(outputs, "IRR")
        payback = self.app._summary_metric(outputs, "Payback Period")

        self.assertIsNotNone(irr)
        self.assertIsNotNone(payback)
        assert irr is not None
        assert payback is not None
        self.assertGreater(irr, 0.0)
        self.assertGreaterEqual(payback, 0.0)

    def test_format_number_returns_na_for_nan(self):
        self.assertEqual(self.app._format_number(float("nan")), "N/A")

    def test_irr_diagnostic_message_explains_missing_sign_change(self):
        cash_flow = self.app.Table(
            [2024, 2025, 2026],
            {self.app.CASH_FLOW_NET_COLUMN: [-3.0, -2.0, -1.0]},
        )
        irr_info = types.SimpleNamespace(
            converged=False,
            value=float("nan"),
            message="Cash flows do not change sign, so IRR is undefined.",
            method="no_sign_change",
        )
        model = types.SimpleNamespace(
            irr_diagnostics=lambda: irr_info,
            cash_flow_statement=lambda: cash_flow,
        )

        message = self.app._irr_diagnostic_message(model)

        self.assertIsNotNone(message)
        assert message is not None
        self.assertIn("Cash flows do not change sign", message)
        self.assertIn("All projected net cash flow periods are non-positive", message)

    def test_irr_diagnostic_message_returns_none_when_irr_available(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(
                encoding="utf-8"
            )
        )
        inputs = self.app.parse_inputs(payload)
        model = self.app.FinancialModel(inputs)
        model.run_core()

        self.assertIsNone(self.app._irr_diagnostic_message(model))

    def test_editor_row_label_combines_name_and_year(self):
        label = self.app._editor_row_label(
            {"Product": "Tablets", "Year": 2027},
            0,
            name_fields=("Product",),
            year_fields=("Year",),
            fallback_prefix="Row",
        )

        self.assertEqual(label, "Tablets | 2027")

    def test_merge_editor_row_updates_preserves_hidden_fields(self):
        merged = self.app._merge_editor_row_updates(
            [{"Product": "Tablets", "__has_fixed__": True, "Fixed Cost": 10.0}],
            [{"Product": "Tablets", "Fixed Cost": 12.0}],
        )

        self.assertEqual(len(merged), 1)
        self.assertTrue(merged[0]["__has_fixed__"])
        self.assertEqual(merged[0]["Fixed Cost"], 12.0)

    def test_projection_horizon_dropdown_updates_years(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(
                encoding="utf-8"
            )
        )

        new_years = list(range(2026, 2031))
        labels = [str(year) for year in new_years]
        self.app._align_payload_horizon(payload, labels, len(new_years), update_years=True)

        self.assertEqual(payload["years"], new_years)
        self.assertEqual(len(payload.get("inflation_series", [])), len(new_years))

        working = payload.get("working_capital", {})
        calendar = working.get("calendar_days", []) if isinstance(working, dict) else []
        self.assertEqual(len(calendar), len(new_years))

    def test_align_payload_preserves_calendar_years_when_not_updating(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(
                encoding="utf-8"
            )
        )

        original_years = list(payload["years"])
        labels = [f"Year {index + 1}" for index in range(len(original_years))]
        self.app._align_payload_horizon(payload, labels, len(labels))

        self.assertEqual(payload["years"], original_years)

    def test_clone_payload_returns_independent_copy(self):
        payload = {"a": {"b": 1}}
        clone = self.app._clone_payload(payload)

        self.assertEqual(clone, payload)
        clone["a"]["b"] = 2
        self.assertEqual(payload["a"]["b"], 1)

    def test_generate_workspace_label_advances_index(self):
        existing = {"Workspace 1": {}, "Workspace 2": {}}
        label = self.app._generate_workspace_label(existing)

        self.assertEqual(label, "Workspace 3")

    def test_core_rows_calculations_match_inputs(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(encoding="utf-8")
        )
        rows = self.app._payload_to_core_rows(payload)
        self.assertTrue(rows)

        inflation = self.app._inflation_factors_from_payload(payload)
        risk = self.app._risk_factors_from_payload(payload)
        years = payload.get("years", [])
        estimates = payload.get("production_estimate", {})

        inflation_factor = inflation[0] if inflation else 1.0
        risk_factor = risk[0] if risk else 1.0

        for row in rows:
            units = float(row["Total Production Units"])
            selling = float(row["Selling Price"])
            production = float(row["Production Cost"])
            freight = float(row["Freight Cost"])
            markup = float(row.get("Markup", 0.0))
            capacity = float(row.get("Max Capacity", 0.0))

            scaled = self.app._scaled_production_series(
                str(row["Product"]), units, years, estimates
            )
            first_year_units = scaled[0] if scaled else 0.0

            expected_revenue = first_year_units * selling * inflation_factor * risk_factor
            expected_cost = (
                first_year_units * (production + freight + markup) * inflation_factor * risk_factor
            )

            self.assertAlmostEqual(row["Total Revenue"], expected_revenue, places=8)
            self.assertAlmostEqual(row["Total Cost"], expected_cost, places=8)
            if capacity > 0:
                self.assertLessEqual(units, capacity + 1e-9)

    def test_core_rows_widget_sync_updates_payload(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(encoding="utf-8")
        )
        rows = self.app._payload_to_core_rows(payload)
        self.assertTrue(rows)

        original = rows[0]
        product_name = str(original["Product"])
        new_price = float(original["Selling Price"]) + 0.25
        new_units = float(original["Total Production Units"]) + 10.0

        self.app.st.session_state[f"core_desc_0"] = product_name
        self.app.st.session_state[f"core_sell_0"] = new_price
        self.app.st.session_state[f"core_units_0"] = new_units

        synced = self.app._sync_core_rows_from_widgets(rows)
        self.assertNotEqual(synced, rows)

        payload_copy = json.loads(json.dumps(payload))
        self.app._core_rows_to_payload(synced, payload_copy)

        parsed = self.app.parse_inputs(payload_copy)
        self.assertAlmostEqual(
            parsed.unit_costs[product_name].selling_price, new_price, places=6
        )

        total_units = sum(parsed.production_estimate[product_name])
        expected_total = float(synced[0]["Total Production Units"]) * len(parsed.years)
        self.assertAlmostEqual(total_units, expected_total, places=6)

    def test_scaled_production_series_treats_total_units_as_annual_capacity(self):
        yearly_units = 123.0
        years = [2024, 2025, 2026]
        prior_profile = {"Tablets": [10.0, 20.0, 30.0]}

        scaled = self.app._scaled_production_series(
            "Tablets",
            yearly_units,
            years,
            prior_profile,
        )

        self.assertEqual(scaled, [yearly_units, yearly_units, yearly_units])

    def test_sync_tax_entries_supports_tax_settings_rows_shape(self):
        rows = [
            {"Year": "2026", "Rate": 0.2},
            {"Year": "2027", "Rate": 0.22},
        ]

        synced = self.app._sync_tax_entries_from_widgets(rows)

        self.assertEqual(
            synced,
            [
                {"label": "2026", "rate": 0.2},
                {"label": "2027", "rate": 0.22},
            ],
        )

    def test_sync_tax_entries_clamps_negative_rates(self):
        rows = [{"Year": "2026", "Rate": -0.1}]

        synced = self.app._sync_tax_entries_from_widgets(rows)

        self.assertEqual(synced, [{"label": "2026", "rate": 0.0}])

    def test_commission_rows_roundtrip(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(encoding="utf-8")
        )
        rows = self.app._payload_to_commission_rows(payload)
        self.assertTrue(rows)

        rows[0]["Yearly Commission %"] = float(rows[0]["Yearly Commission %"]) + 1.5
        rows[0]["Revenue Share (%)"] = 80.0
        rows[0]["Revenue"] = None
        rows[0]["Payment Days"] = int(rows[0]["Payment Days"]) + 10

        self.app._commission_rows_to_payload(rows, payload)

        commission_section = payload.get("distributor_commission", {})
        stored_rows = commission_section.get("rows", [])
        self.assertTrue(stored_rows)
        first = stored_rows[0]
        product = first["product"]
        target_year = first["year"]
        base_rates = self.app._commission_base_rates(payload)
        rate = base_rates.get(product, 0.05)
        product_rows = [
            row for row in rows if row.get("Product") == product
        ]
        product_rows.sort(key=lambda row: row.get("Year", 0))
        expected_rate = rate
        for entry in product_rows:
            increment = float(entry.get("Yearly Commission %", 0.0))
            expected_rate = expected_rate * (1 + increment / 100.0)
            if entry.get("Year") == target_year:
                break
        self.assertAlmostEqual(first["rate"], expected_rate, delta=0.002)
        self.assertAlmostEqual(first["revenue_share"] * 100.0, rows[0]["Revenue Share (%)"], places=6)
        self.assertEqual(first["payment_days"], rows[0]["Payment Days"])

    def test_utility_entries_extend_projection_horizon(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(encoding="utf-8")
        )
        rows = self.app._payload_to_utility_entries(payload)
        original_years = list(payload["years"])
        new_year = original_years[-1] + 1

        new_row = dict(rows[-1]) if rows else self.app._default_utility_entry(0)
        new_row["label"] = str(new_year)
        new_row["year"] = new_year
        rows.append(new_row)

        self.app._utility_entries_to_payload(rows, payload)

        stored_rows = payload.get("utility_costs", {}).get("years", [])
        self.assertEqual(len(stored_rows), len(rows))
        self.assertEqual(int(stored_rows[-1]["year"]), new_year)

        self.app._align_payload_horizon(
            payload,
            [str(row.get("label", "")) for row in rows],
            len(rows),
            update_years=True,
        )

        self.assertEqual(payload["years"][-1], new_year)

    def test_receivable_rows_do_not_shrink_horizon(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(encoding="utf-8")
        )
        original_years = list(payload["years"])
        rows = self.app._payload_to_receivable_rows(payload)

        trimmed = rows[:2]
        self.app._receivable_rows_to_payload(trimmed, payload)

        self.assertEqual(payload["years"], original_years)

    def test_cached_parse_and_model_run_reuse_digest(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(encoding="utf-8")
        )
        inputs_a, digest_a = self.app._cached_parse_inputs(payload)
        inputs_b, digest_b = self.app._cached_parse_inputs(payload)

        self.assertIs(inputs_a, inputs_b)
        self.assertEqual(digest_a, digest_b)

        model1, outputs1 = self.app._cached_model_run(inputs_a, digest_a)
        model2, outputs2 = self.app._cached_model_run(inputs_b, digest_b)

        self.assertIs(model1, model2)
        self.assertIs(outputs1, outputs2)

    def test_inventory_rows_roundtrip(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(encoding="utf-8")
        )
        rows = self.app._payload_to_inventory_rows(payload)
        self.assertTrue(rows)

        rows[0]["inventory_days"] = float(rows[0]["inventory_days"]) + 5
        rows[0]["accounts_payable_days"] = float(rows[0]["accounts_payable_days"]) + 3
        rows[0]["days_in_year"] = float(rows[0]["days_in_year"]) - 1

        self.app._inventory_rows_to_payload(rows, payload)

        working = payload.get("working_capital", {})
        calendar = working.get("calendar_days", [])
        day_mapping = working.get("days", {})
        inventory = day_mapping.get("inventory", [])
        payables = day_mapping.get("accounts_payable", [])

        self.assertAlmostEqual(calendar[0], rows[0]["days_in_year"], places=6)
        self.assertAlmostEqual(inventory[0], rows[0]["inventory_days"], places=6)
        self.assertAlmostEqual(payables[0], rows[0]["accounts_payable_days"], places=6)

    def test_receivable_rows_roundtrip(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(encoding="utf-8")
        )
        rows = self.app._payload_to_receivable_rows(payload)
        self.assertTrue(rows)

        rows[0]["accounts_receivable_days"] = float(rows[0]["accounts_receivable_days"]) + 4
        rows[0]["prepaid_expense_days"] = float(rows[0]["prepaid_expense_days"]) + 2
        rows[0]["other_asset_days"] = float(rows[0]["other_asset_days"]) + 1
        rows[0]["days_in_year"] = float(rows[0]["days_in_year"]) + 2

        self.app._receivable_rows_to_payload(rows, payload)

        working = payload.get("working_capital", {})
        calendar = working.get("calendar_days", [])
        day_mapping = working.get("days", {})
        receivable = day_mapping.get("accounts_receivable", [])
        prepaid = day_mapping.get("prepaid_expenses", [])
        other_assets = day_mapping.get("other_assets", [])

        self.assertAlmostEqual(calendar[0], rows[0]["days_in_year"], places=6)
        self.assertAlmostEqual(receivable[0], rows[0]["accounts_receivable_days"], places=6)
        self.assertAlmostEqual(prepaid[0], rows[0]["prepaid_expense_days"], places=6)
        self.assertAlmostEqual(other_assets[0], rows[0]["other_asset_days"], places=6)

    def test_ai_settings_roundtrip(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(encoding="utf-8")
        )
        settings = self.app._payload_to_ai_settings(payload)
        settings["provider"] = "Azure OpenAI"
        settings["model"] = "gpt-custom"
        settings["ml_methods"] = ["linear_regression", "moving_average"]
        settings["generative_features"] = ["summary", "risk_review"]
        settings["api_key"] = "test-key"
        self.app._ai_settings_to_payload(settings, payload)
        updated = self.app._payload_to_ai_settings(payload)
        self.assertEqual(updated["provider"], "Azure OpenAI")
        self.assertEqual(updated["model"], "gpt-custom")
        self.assertIn("moving_average", updated["ml_methods"])
        self.assertIn("risk_review", updated["generative_features"])
        self.assertEqual(updated["api_key"], "test-key")

    def test_scenario_payload_uses_configured_adjustments(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(encoding="utf-8")
        )
        inputs, digest = self.app._cached_parse_inputs(payload)
        base_model, base_outputs = self.app._cached_model_run(inputs, digest)

        snapshot = self.app._clone_payload(payload)
        scenario_model, scenario_outputs = self.app._ensure_scenario_payload(
            "best", snapshot, base_model, base_outputs
        )

        self.assertIsNot(scenario_model, base_model)
        best_config = payload["scenarios"]["best"]
        self.assertAlmostEqual(
            scenario_model.inputs.financing.discount_rate, best_config["interest"][0]
        )
        self.assertIsNotNone(scenario_outputs)

        same_model, same_outputs = self.app._ensure_scenario_payload(
            "unknown", snapshot, base_model, base_outputs
        )
        self.assertIs(same_model, base_model)
        self.assertIs(same_outputs, base_outputs)

    def test_generate_excel_bytes_returns_workbook(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(encoding="utf-8")
        )
        inputs, digest = self.app._cached_parse_inputs(payload)
        model, outputs = self.app._cached_model_run(inputs, digest)

        workbook = self.app._generate_excel_bytes(model, outputs, "Base")
        self.assertIsInstance(workbook, (bytes, bytearray))
        self.assertGreater(len(workbook), 0)

    def test_labor_model_payload_helpers_roundtrip(self):
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(encoding="utf-8")
        )
        years = payload["years"]
        payload.setdefault("labor", {})["model_v1"] = {
            "roles": [
                {
                    "name": "Operators",
                    "labor_type": "direct",
                    "behavior": "variable",
                    "headcount": 5,
                    "salary": 1.0,
                    "planned_headcount": [5 for _ in years],
                    "benefits_rate": [0.1 for _ in years],
                    "overtime_rate": [0.05 for _ in years],
                    "burden_rate": [0.2 for _ in years],
                    "productivity_target": [100000 for _ in years],
                    "source": "HR",
                    "owner": "Ops",
                    "benchmark_year": "2024",
                }
            ],
            "settings": {
                "shifts": [1 for _ in years],
                "utilization": [0.85 for _ in years],
                "operating_hours_per_shift": [2080 for _ in years],
                "absenteeism": [0.03 for _ in years],
                "overtime_cap": [0.1 for _ in years],
                "hiring_delay_quarters": [1 for _ in years],
                "contractor_hours": [0 for _ in years],
                "contractor_rate": [0 for _ in years],
                "transition_training_cost": [0 for _ in years],
                "supervision_increment": [0 for _ in years],
                "shift_allowance": [0 for _ in years],
                "wage_escalation_direct": [0.03 for _ in years],
                "wage_escalation_indirect": [0.02 for _ in years],
            },
        }

        role_rows = self.app._payload_to_labor_model_rows(payload)
        settings_rows = self.app._payload_to_labor_model_settings_rows(payload)
        self.assertTrue(role_rows)
        self.assertEqual(len(settings_rows), len(years))

        role_rows[0]["Name"] = "Operators Revised"
        role_rows[0]["Planned Headcount"] = "5, 6, 7"
        settings_rows[0]["Shifts"] = 2.0

        updated_payload = json.loads(json.dumps(payload))
        self.app._labor_model_rows_to_payload(role_rows, updated_payload)
        self.app._labor_model_settings_rows_to_payload(settings_rows, updated_payload)

        model_v1 = updated_payload["labor"]["model_v1"]
        self.assertEqual(model_v1["roles"][0]["name"], "Operators Revised")
        self.assertEqual(model_v1["roles"][0]["planned_headcount"][1], 6.0)
        self.assertEqual(model_v1["settings"]["shifts"][0], 2.0)

    def test_build_monte_correlation_matrix_rows_uses_upper_triangle(self):
        variable_map = {
            "revenue_growth": "Revenue Growth",
            "raw_material_cost": "Cost of Materials",
            "labor_cost": "Labour",
        }
        rows = self.app._build_monte_correlation_matrix_rows(
            ["revenue_growth", "raw_material_cost", "labor_cost"],
            variable_map,
            {
                "revenue_growth": {
                    "raw_material_cost": -0.2,
                    "labor_cost": -0.1,
                },
                "labor_cost": {
                    "raw_material_cost": 0.25,
                },
            },
        )

        self.assertEqual(rows[0]["Revenue Growth"], 1.0)
        self.assertEqual(rows[0]["Cost of Materials"], -0.2)
        self.assertEqual(rows[0]["Labour"], -0.1)
        self.assertIsNone(rows[1]["Revenue Growth"])
        self.assertEqual(rows[1]["Labour"], 0.25)
        self.assertIsNone(rows[2]["Revenue Growth"])
        self.assertIsNone(rows[2]["Cost of Materials"])
        self.assertEqual(rows[2]["Labour"], 1.0)

    def test_merge_monte_correlation_matrix_rows_preserves_non_table_entries(self):
        variable_map = {
            "revenue_growth": "Revenue Growth",
            "raw_material_cost": "Cost of Materials",
            "labor_cost": "Labour",
        }
        merged = self.app._merge_monte_correlation_matrix_rows(
            [
                {
                    "Variable": "Revenue Growth",
                    "Revenue Growth": 1.0,
                    "Cost of Materials": -0.35,
                    "Labour": -0.1,
                },
                {
                    "Variable": "Cost of Materials",
                    "Revenue Growth": None,
                    "Cost of Materials": 1.0,
                    "Labour": 0.2,
                },
                {
                    "Variable": "Labour",
                    "Revenue Growth": None,
                    "Cost of Materials": None,
                    "Labour": 1.0,
                },
            ],
            ["revenue_growth", "raw_material_cost", "labor_cost"],
            variable_map,
            {
                "revenue_growth": {
                    "selling_price": 0.15,
                    "raw_material_cost": -0.2,
                },
                "other": {
                    "revenue_growth": 0.05,
                },
            },
        )

        self.assertEqual(merged["revenue_growth"]["selling_price"], 0.15)
        self.assertEqual(merged["other"]["revenue_growth"], 0.05)
        self.assertEqual(merged["revenue_growth"]["raw_material_cost"], -0.35)
        self.assertEqual(merged["raw_material_cost"]["revenue_growth"], -0.35)
        self.assertEqual(merged["revenue_growth"]["labor_cost"], -0.1)
        self.assertEqual(merged["labor_cost"]["revenue_growth"], -0.1)
        self.assertEqual(merged["raw_material_cost"]["labor_cost"], 0.2)
        self.assertEqual(merged["labor_cost"]["raw_material_cost"], 0.2)


class UploadLoaderTests(unittest.TestCase):
    def setUp(self):
        if "pharma_financial.app" in sys.modules:
            importlib.reload(sys.modules["pharma_financial.app"])
            self.app = sys.modules["pharma_financial.app"]
        else:
            self.app = importlib.import_module("pharma_financial.app")

        self.sample_payload = {
            "example": True,
            "years": [2024, 2025],
        }
        self.json_text = json.dumps(self.sample_payload)

    def test_load_from_json_bytes(self):
        loaded = self.app._load_payload_from_bytes(self.json_text.encode("utf-8"), ".json")
        self.assertEqual(loaded["example"], True)

    def test_load_from_csv_fragment(self):
        csv_text = f"header,value\njson,{self.json_text}\n"
        loaded = self.app._load_payload_from_bytes(csv_text.encode("utf-8"), ".csv")
        self.assertEqual(loaded["years"], [2024, 2025])

    @unittest.skipUnless(Workbook is not None, "openpyxl is required for this test")
    def test_load_from_excel_fragment(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet["A1"] = "Assumptions"
        sheet["A2"] = self.json_text
        buffer = BytesIO()
        workbook.save(buffer)
        loaded = self.app._load_payload_from_bytes(buffer.getvalue(), ".xlsx")
        self.assertIn("example", loaded)

    @unittest.skipUnless(Document is not None, "python-docx is required for this test")
    def test_load_from_docx_fragment(self):
        document = Document()
        document.add_paragraph("Example assumptions")
        document.add_paragraph(self.json_text)
        buffer = BytesIO()
        document.save(buffer)
        loaded = self.app._load_payload_from_bytes(buffer.getvalue(), ".docx")
        self.assertEqual(loaded["years"], [2024, 2025])

    @unittest.skipUnless(FPDF is not None, "fpdf is required for this test")
    def test_load_from_pdf_fragment(self):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, txt=f"Example assumptions\n{self.json_text}")
        pdf_bytes = pdf.output(dest="S").encode("latin-1")
        loaded = self.app._load_payload_from_bytes(pdf_bytes, ".pdf")
        self.assertTrue(loaded["example"])


class SetupTabRenderTests(unittest.TestCase):
    def setUp(self):
        self.setup_tab = importlib.import_module("pharma_financial.ui.tabs.setup")
        self.app = importlib.import_module("pharma_financial.app")
        self.original_streamlit = self.app.st
        self.original_header = self.setup_tab.shell.render_section_header
        self.original_collapsible = self.setup_tab._render_collapsible_section
        self.original_table = self.setup_tab._render_table_like
        payload = json.loads(
            Path("src/pharma_financial/data/default_inputs.json").read_text(
                encoding="utf-8"
            )
        )

        fake_streamlit = types.SimpleNamespace(
            session_state={"input_payload": payload},
            markdown=lambda *args, **kwargs: None,
        )
        self.app.st = fake_streamlit

        self.setup_tab.shell.render_section_header = lambda *args, **kwargs: None
        self.setup_tab._render_collapsible_section = lambda *args, **kwargs: None
        self.setup_tab._render_table_like = lambda *args, **kwargs: None

    def tearDown(self):
        self.app.st = self.original_streamlit
        self.setup_tab.shell.render_section_header = self.original_header
        self.setup_tab._render_collapsible_section = self.original_collapsible
        self.setup_tab._render_table_like = self.original_table

    def test_render_commercial_operations_keeps_outputs_available(self):
        outputs = types.SimpleNamespace(
            commercial_diagnostics=[{"metric": "Gross Margin", "value": "68%"}]
        )

        self.setup_tab.render_commercial_operations(None, None, outputs, "digest")

    def test_render_funding_working_capital_keeps_outputs_available(self):
        outputs = types.SimpleNamespace(
            sources_and_uses=[{"source": "Equity", "amount": 100.0}],
            covenant_headroom=[{"year": 2026, "headroom": 1.4}],
        )

        self.setup_tab.render_funding_working_capital(None, None, outputs, "digest")


if __name__ == "__main__":  # pragma: no cover
    unittest.main()
