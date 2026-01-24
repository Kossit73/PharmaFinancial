"""Utilities for assembling and exporting consolidated model reports."""
from __future__ import annotations

import json
import math
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO, StringIO
from typing import Any, Iterable, List, Mapping, MutableMapping, Sequence, Tuple
from xml.sax.saxutils import escape as xml_escape
from zipfile import ZipFile

from .model import FinancialModel, FinancialOutputs
from .table import Table

try:  # pragma: no cover - optional dependency
    import pandas as pd  # type: ignore
except Exception:  # pragma: no cover - pandas optional in minimal environments
    pd = None  # type: ignore

try:  # pragma: no cover - optional dependency
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - python-docx may not be installed
    Document = None  # type: ignore

try:  # pragma: no cover - optional dependency
    from fpdf import FPDF  # type: ignore
except Exception:  # pragma: no cover - FPDF may not be installed
    FPDF = None  # type: ignore

try:  # pragma: no cover - optional dependency for Excel styling
    from openpyxl.styles import Alignment, Font, PatternFill  # type: ignore
    from openpyxl.utils import get_column_letter  # type: ignore
except Exception:  # pragma: no cover - optional dependency for styling
    Alignment = None  # type: ignore
    Font = None  # type: ignore
    PatternFill = None  # type: ignore
    get_column_letter = None  # type: ignore

JSON_MIME = "application/json"
CSV_MIME = "text/csv"
EXCEL_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
WORD_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"

REPORT_FORMATS = ["PDF", "Word", "Excel", "CSV", "JSON"]


class ReportGenerationError(RuntimeError):
    """Raised when the requested report cannot be produced."""


@dataclass
class ReportTable:
    """Container describing an individual table within a report."""

    title: str
    data: Any
    note: str | None = None


@dataclass
class ReportSection:
    """Represents a logical report section grouping multiple tables."""

    title: str
    tables: List[ReportTable]
    notes: List[str] | None = None


def collect_report_sections(model: FinancialModel, outputs: FinancialOutputs) -> List[ReportSection]:
    """Assemble report sections spanning all dashboards in the required order."""

    sections: List[ReportSection] = []

    key_tables: List[ReportTable] = [
        ReportTable("Summary Metrics", outputs.summary_metrics),
        ReportTable("Goal Seek", outputs.goal_seek),
    ]

    try:
        key_tables.append(ReportTable("Working Capital Schedule", model.working_capital_schedule()))
    except Exception as exc:  # pragma: no cover - defensive guard for runtime calculations
        key_tables.append(ReportTable("Working Capital Schedule", [], note=f"Unavailable: {exc}"))

    try:
        key_tables.append(ReportTable("Inventory Schedule", model.inventory_schedule()))
    except Exception as exc:  # pragma: no cover - defensive guard for runtime calculations
        key_tables.append(ReportTable("Inventory Schedule", [], note=f"Unavailable: {exc}"))

    ai_notes: List[str] = []
    if outputs.ai_insights is not None:
        insight = outputs.ai_insights
        if insight.ml_forecast is not None:
            key_tables.append(ReportTable("AI Revenue Forecast", insight.ml_forecast))
        if insight.generative_summary:
            ai_notes.append(insight.generative_summary)
        if insight.metadata:
            provider = insight.metadata.get("provider")
            model_name = insight.metadata.get("model")
            status = insight.metadata.get("status")
            details = ", ".join(
                str(value)
                for value in [provider, model_name, status]
                if value not in (None, "")
            )
            if details:
                ai_notes.append(f"AI configuration: {details}")

    sections.append(ReportSection("Key Metrics Dashboard", key_tables, notes=ai_notes or None))

    perf_tables: List[ReportTable] = [
        ReportTable(
            "Statement of Financial Performance",
            outputs.income_statement.rounded(0, exclude_keywords=("Margin", "Return")),
        ),
    ]
    try:
        perf_tables.append(ReportTable("Gross Revenue Schedule", model.revenue_schedule()))
    except Exception as exc:  # pragma: no cover - defensive guard
        perf_tables.append(ReportTable("Gross Revenue Schedule", [], note=f"Unavailable: {exc}"))

    try:
        perf_tables.append(ReportTable("Total Expenses Schedule", model.cost_structure()))
    except Exception as exc:  # pragma: no cover - defensive guard
        perf_tables.append(ReportTable("Total Expenses Schedule", [], note=f"Unavailable: {exc}"))

    sections.append(ReportSection("Financial Performance", perf_tables))

    break_even_tables = [
        ReportTable("Break-even Analysis", outputs.break_even),
        ReportTable("Payback Schedule", outputs.payback),
        ReportTable("Discounted Payback Schedule", outputs.discounted_payback),
    ]
    sections.append(ReportSection("Break-even & Payback", break_even_tables))

    sections.append(
        ReportSection(
            "Financial Position",
            [
                ReportTable(
                    "Statement of Financial Position",
                    outputs.balance_sheet.rounded(0),
                )
            ],
        )
    )

    sections.append(
        ReportSection(
            "Cash Flow Statement",
            [
                ReportTable(
                    "Statement of Cash Flows",
                    outputs.cash_flow.rounded(0),
                )
            ],
        )
    )

    sensitivity_tables: List[ReportTable] = []
    for name, table in outputs.sensitivity_results.items():
        label = name.replace("_", " ").title()
        sensitivity_tables.append(ReportTable(f"Sensitivity: {label}", table))
    if not sensitivity_tables:
        sensitivity_tables.append(ReportTable("Sensitivity Analysis", [], note="No sensitivity configurations provided."))
    sections.append(ReportSection("Sensitivity Analysis", sensitivity_tables))

    scenario_tables: List[ReportTable] = []
    for name, table in outputs.scenario_results.items():
        scenario_tables.append(ReportTable(f"Scenario: {name}", table))
    for key, result in outputs.scenario_tool_results.items():
        label = key.replace("_", " ").title()
        scenario_tables.append(ReportTable(f"Scenario Tool: {label}", result.rows, note=result.interpretation))
    if not scenario_tables:
        scenario_tables.append(ReportTable("Scenario / IFs Analysis", [], note="No scenario configurations provided."))
    sections.append(ReportSection("Scenario / IFs Analysis", scenario_tables))

    sections.append(ReportSection("Monte Carlo Simulation", [ReportTable("Monte Carlo Simulation", outputs.monte_carlo)]))

    return sections


def generate_report(
    sections: Sequence[ReportSection],
    format_name: str,
    *,
    report_name: str = "longevity_financial_report",
) -> Tuple[bytes, str, str]:
    """Generate a consolidated report in the requested format."""

    if not sections:
        raise ReportGenerationError("No sections available to generate a report.")

    fmt = format_name.strip().lower()
    if fmt not in {f.lower() for f in REPORT_FORMATS}:
        raise ReportGenerationError(f"Unsupported report format: {format_name}")

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")

    if fmt == "json":
        data = _build_json(sections)
        filename = f"{report_name}_{timestamp}.json"
        return data, JSON_MIME, filename

    if fmt == "csv":
        data = _build_csv(sections)
        filename = f"{report_name}_{timestamp}.csv"
        return data, CSV_MIME, filename

    if fmt == "excel":
        data = _build_excel(sections)
        filename = f"{report_name}_{timestamp}.xlsx"
        return data, EXCEL_MIME, filename

    if fmt == "word":
        data = _build_word(sections)
        filename = f"{report_name}_{timestamp}.docx"
        return data, WORD_MIME, filename

    if fmt == "pdf":
        data = _build_pdf(sections)
        filename = f"{report_name}_{timestamp}.pdf"
        return data, PDF_MIME, filename

    raise ReportGenerationError(f"Unsupported report format: {format_name}")


# ---------------------------------------------------------------------------
# format builders
# ---------------------------------------------------------------------------

def _build_json(sections: Sequence[ReportSection]) -> bytes:
    payload: List[MutableMapping[str, Any]] = []
    for section in sections:
        entry: MutableMapping[str, Any] = {
            "title": section.title,
            "tables": [],
        }
        if section.notes:
            entry["notes"] = list(section.notes)
        for table in section.tables:
            entry["tables"].append(
                {
                    "title": table.title,
                    "rows": _table_to_rows(table.data),
                    **({"note": table.note} if table.note else {}),
                }
            )
        payload.append(entry)
    return json.dumps(
        {"generated": datetime.now(timezone.utc).isoformat(), "sections": payload},
        indent=2,
    ).encode("utf-8")


def _build_csv(sections: Sequence[ReportSection]) -> bytes:
    buffer = StringIO()
    for section in sections:
        buffer.write(f"# Section: {section.title}\n")
        if section.notes:
            for note in section.notes:
                buffer.write(f"# Note: {note}\n")
        for table in section.tables:
            buffer.write(f"## Table: {table.title}\n")
            if table.note:
                buffer.write(f"## Note: {table.note}\n")
            rows = _table_to_rows(table.data)
            if not rows:
                buffer.write("(no data)\n\n")
                continue
            columns = list(rows[0].keys())
            buffer.write(",".join(_escape_csv(value) for value in columns) + "\n")
            for row in rows:
                buffer.write(",".join(_escape_csv(row.get(column, "")) for column in columns) + "\n")
            buffer.write("\n")
        buffer.write("\n")
    return buffer.getvalue().encode("utf-8")


def _build_excel(sections: Sequence[ReportSection]) -> bytes:
    entries = _collect_sheet_entries(sections)

    if pd is not None:
        buffer = BytesIO()
        try:
            with pd.ExcelWriter(buffer) as writer:  # type: ignore[arg-type]
                for entry in entries:
                    frame = _rows_to_dataframe(entry["rows"])
                    if frame is None:
                        raise ValueError("pandas unavailable")
                    frame = frame.copy()
                    note = entry.get("note")
                    if note:
                        note_col = "Notes"
                        if frame.empty:
                            frame = pd.DataFrame({note_col: [note]})
                        else:
                            if note_col not in frame.columns:
                                frame[note_col] = ""
                            frame.iloc[0, frame.columns.get_loc(note_col)] = note
                    frame.to_excel(writer, sheet_name=entry["sheet_name"], index=False)
                    sheet = writer.sheets.get(entry["sheet_name"])
                    if sheet and Font and PatternFill and Alignment and get_column_letter:
                        header_fill = PatternFill(fill_type="solid", fgColor="1F4E78")
                        header_font = Font(color="FFFFFF", bold=True)
                        for cell in sheet[1]:
                            cell.font = header_font
                            cell.fill = header_fill
                            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                        sheet.freeze_panes = "A2"
                        for idx, column_cells in enumerate(sheet.columns, start=1):
                            max_length = 0
                            for cell in column_cells[: min(len(column_cells), 50)]:
                                value = "" if cell.value is None else str(cell.value)
                                max_length = max(max_length, len(value))
                            adjusted = min(max_length + 2, 60)
                            sheet.column_dimensions[get_column_letter(idx)].width = adjusted
            return buffer.getvalue()
        except Exception:  # pragma: no cover - fall back to pure-Python writer
            pass

    return _build_excel_fallback(entries)


def _build_word(sections: Sequence[ReportSection]) -> bytes:
    if Document is not None:
        document = Document()
        document.add_heading("Pharmaceuticals Financial Report", level=1)
        document.add_paragraph(
            f"Generated on {datetime.now(timezone.utc).isoformat()} UTC",
            style="Subtitle",
        )
        for section in sections:
            document.add_heading(section.title, level=2)
            if section.notes:
                for note in section.notes:
                    document.add_paragraph(note, style="List Bullet")
            for table in section.tables:
                document.add_heading(table.title, level=3)
                if table.note:
                    document.add_paragraph(table.note, style="List Bullet")
                rows = _table_to_rows(table.data)
                if not rows:
                    document.add_paragraph("No data available.")
                    continue
                columns = _ordered_columns(rows)
                if not columns:
                    document.add_paragraph("No data available.")
                    continue
                word_table = document.add_table(rows=1, cols=len(columns))
                word_table.style = "Table Grid"
                header_cells = word_table.rows[0].cells
                for idx, column in enumerate(columns):
                    header_cells[idx].text = str(column)
                for row in rows:
                    row_cells = word_table.add_row().cells
                    for idx, column in enumerate(columns):
                        row_cells[idx].text = _format_report_value(row.get(column, ""))
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    return _build_word_fallback(_collect_report_blocks(sections))


def _build_pdf(sections: Sequence[ReportSection]) -> bytes:
    if FPDF is not None:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, _safe_pdf_text("Pharmaceuticals Financial Report"), ln=True)
        pdf.set_font("Arial", "", 10)
        pdf.cell(
            0,
            8,
            _safe_pdf_text(f"Generated on {datetime.now(timezone.utc).isoformat()} UTC"),
            ln=True,
        )
        for section in sections:
            pdf.ln(4)
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, _safe_pdf_text(section.title.upper()), ln=True)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            if section.notes:
                pdf.set_font("Arial", "", 10)
                for note in section.notes:
                    _pdf_multiline(pdf, f"- {_safe_pdf_text(note)}")
            for table in section.tables:
                pdf.ln(2)
                pdf.set_font("Arial", "B", 11)
                pdf.cell(0, 7, _safe_pdf_text(table.title), ln=True)
                if table.note:
                    pdf.set_font("Arial", "", 9)
                    _pdf_multiline(pdf, f"Note: {_safe_pdf_text(table.note)}")
                rows = _table_to_rows(table.data)
                if not rows:
                    pdf.set_font("Arial", "", 9)
                    _pdf_multiline(pdf, "No data available.")
                    continue
                columns = _ordered_columns(rows)
                if not columns:
                    pdf.set_font("Arial", "", 9)
                    _pdf_multiline(pdf, "No data available.")
                    continue
                _render_pdf_table(pdf, columns, rows)
        return bytes(pdf.output(dest="S").encode("latin-1"))

    return _build_pdf_fallback(_collect_report_blocks(sections))


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------


def _collect_report_blocks(sections: Sequence[ReportSection]) -> List[Tuple[str, str]]:
    """Flatten sections into typed text blocks for downstream exporters."""

    blocks: List[Tuple[str, str]] = [
        ("title", "Pharmaceuticals Financial Report"),
        ("subtitle", f"Generated on {datetime.now(timezone.utc).isoformat()} UTC"),
    ]

    for section in sections:
        blocks.append(("heading1", section.title))
        if section.notes:
            for note in section.notes:
                blocks.append(("note", f"Note: {note}"))
        for table in section.tables:
            blocks.append(("heading2", table.title))
            if table.note:
                blocks.append(("note", f"Note: {table.note}"))
            rows = _table_to_rows(table.data)
            if not rows:
                blocks.append(("body", "No data available."))
                continue
            columns = _ordered_columns(rows)
            if columns:
                blocks.append(("body", " | ".join(columns)))
                for row in rows:
                    values = [_stringify(row.get(column, "")) for column in columns]
                    blocks.append(("body", " | ".join(values)))
            else:
                for row in rows:
                    values = [_stringify(value) for value in row.values()]
                    blocks.append(("body", " | ".join(values) or "No data available."))
            blocks.append(("body", ""))

    return blocks


def _safe_pdf_text(text: str) -> str:
    """Return text safe for FPDF's Latin-1 encoding."""

    try:
        text.encode("latin-1")
        return text
    except UnicodeEncodeError:
        return text.encode("latin-1", "replace").decode("latin-1")


def _format_report_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        if math.isnan(value):
            return ""
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        absolute = abs(float(value))
        formatted = f"{absolute:,.4f}" if absolute < 1 else f"{absolute:,.2f}"
        return f"({formatted})" if value < 0 else formatted
    return str(value)


def _truncate_text(pdf: "FPDF", text: str, width: float) -> str:
    if pdf.get_string_width(text) <= width:
        return text
    trimmed = text
    while trimmed and pdf.get_string_width(f"{trimmed}…") > width:
        trimmed = trimmed[:-1]
    return f"{trimmed}…" if trimmed else text[:1]


def _render_pdf_table(
    pdf: "FPDF",
    columns: Sequence[str],
    rows: Sequence[Mapping[str, Any]],
) -> None:
    page_width = pdf.w - pdf.l_margin - pdf.r_margin
    if not columns:
        return
    first_width = page_width * 0.45 if len(columns) > 1 else page_width
    other_width = (
        (page_width - first_width) / max(len(columns) - 1, 1) if len(columns) > 1 else 0
    )
    widths = [first_width] + [other_width] * (len(columns) - 1)

    pdf.set_font("Arial", "B", 9)
    pdf.set_fill_color(230, 230, 230)
    for idx, column in enumerate(columns):
        text = _truncate_text(pdf, _safe_pdf_text(str(column)), widths[idx])
        align = "L" if idx == 0 else "R"
        pdf.cell(widths[idx], 7, text, border="B", align=align, fill=True)
    pdf.ln(7)

    pdf.set_font("Arial", "", 9)
    for row in rows:
        for idx, column in enumerate(columns):
            raw = row.get(column, "")
            value = _format_report_value(raw)
            align = "L" if idx == 0 else "R"
            text = _truncate_text(pdf, _safe_pdf_text(value), widths[idx])
            pdf.cell(widths[idx], 6, text, border="B", align=align)
        pdf.ln(6)


def _collect_sheet_entries(sections: Sequence[ReportSection]) -> List[dict[str, Any]]:
    """Prepare sheet metadata shared by the Excel exporters."""

    tracker: dict[str, int] = {}
    entries: List[dict[str, Any]] = []

    for section in sections:
        for table in section.tables:
            rows = _table_to_rows(table.data)
            sheet_name = _sheet_name(section.title, table.title, tracker)
            entries.append({"sheet_name": sheet_name, "rows": rows, "note": table.note})
        if section.notes:
            sheet_name = _sheet_name(section.title, "Notes", tracker)
            note_rows = [{"Notes": note} for note in section.notes]
            entries.append({"sheet_name": sheet_name, "rows": note_rows, "note": None})

    if not entries:
        entries.append(
            {
                "sheet_name": "Report",
                "rows": [{"Message": "No data available."}],
                "note": None,
            }
        )

    return entries


def _rows_to_dataframe(rows: List[Mapping[str, Any]]):
    """Convert a list of dictionaries to a pandas DataFrame when available."""

    if pd is None:
        return None
    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


def _build_excel_fallback(entries: Sequence[Mapping[str, Any]]) -> bytes:
    """Create a minimal XLSX archive without third-party dependencies."""

    sheets: List[Tuple[str, List[List[str]]]] = []
    for entry in entries:
        rows = entry.get("rows", [])
        note = entry.get("note")
        matrix = _rows_to_matrix(rows, note)
        sheets.append((entry["sheet_name"], matrix))

    return _write_xlsx_archive(sheets)


def _rows_to_matrix(rows: Sequence[Mapping[str, Any]], note: Any) -> List[List[str]]:
    """Normalise rows into a rectangular matrix of strings."""

    columns = _ordered_columns(rows)
    matrix: List[List[str]] = []

    if columns:
        matrix.append(columns)
        for row in rows:
            matrix.append([_stringify(row.get(column, "")) for column in columns])
    else:
        matrix.append(["Value"])
        if rows:
            for row in rows:
                value = next(iter(row.values()), "")
                matrix.append([_stringify(value)])
        else:
            matrix.append(["No data available."])
            note = None

    if note:
        padding = [""] * (len(matrix[0]) - 1)
        matrix.append([f"Note: {note}", *padding])

    # Ensure rectangular shape
    width = max(len(row) for row in matrix)
    for row in matrix:
        if len(row) < width:
            row.extend([""] * (width - len(row)))

    return matrix


def _write_xlsx_archive(sheets: Sequence[Tuple[str, List[List[str]]]]) -> bytes:
    buffer = BytesIO()
    sheet_count = max(len(sheets), 1)
    sheet_names = [name for name, _ in sheets] or ["Report"]

    with ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", _excel_content_types(sheet_count))
        archive.writestr("_rels/.rels", _excel_root_rels())
        archive.writestr("docProps/core.xml", _excel_core_props())
        archive.writestr("docProps/app.xml", _excel_app_props(sheet_names))
        archive.writestr("xl/_rels/workbook.xml.rels", _excel_workbook_rels(sheet_count))
        archive.writestr("xl/workbook.xml", _excel_workbook_xml(sheet_names))
        archive.writestr("xl/styles.xml", _excel_styles_xml())

        if sheets:
            for index, (_, rows) in enumerate(sheets, start=1):
                archive.writestr(f"xl/worksheets/sheet{index}.xml", _excel_sheet_xml(rows))
        else:
            archive.writestr("xl/worksheets/sheet1.xml", _excel_sheet_xml([["Message"], ["No data available."]]))

    return buffer.getvalue()


def _excel_sheet_xml(rows: Sequence[Sequence[str]]) -> str:
    max_cols = max(len(row) for row in rows) if rows else 1
    parts = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">',
        "<sheetData>",
    ]

    for r_idx, row in enumerate(rows, start=1):
        parts.append(f'<row r="{r_idx}">')
        padded = list(row) + [""] * (max_cols - len(row))
        for c_idx, value in enumerate(padded, start=1):
            column = _excel_column_letter(c_idx)
            escaped = xml_escape(_stringify(value))
            parts.append(
                f'<c r="{column}{r_idx}" t="inlineStr"><is><t xml:space="preserve">{escaped}</t></is></c>'
            )
        parts.append("</row>")

    parts.append("</sheetData></worksheet>")
    return "\n".join(parts)


def _excel_column_letter(index: int) -> str:
    result = ""
    while index > 0:
        index, remainder = divmod(index - 1, 26)
        result = chr(65 + remainder) + result
    return result or "A"


def _excel_content_types(count: int) -> str:
    sheet_overrides = "".join(
        f'<Override PartName="/xl/worksheets/sheet{i}.xml" '
        f'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        for i in range(1, count + 1)
    )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        '<Override PartName="/xl/styles.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>'
        '<Override PartName="/docProps/app.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
        f"{sheet_overrides}</Types>"
    )


def _excel_root_rels() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="xl/workbook.xml"/>'
        '<Relationship Id="rId2" '
        'Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" '
        'Target="docProps/core.xml"/>'
        '<Relationship Id="rId3" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" '
        'Target="docProps/app.xml"/>'
        "</Relationships>"
    )


def _excel_workbook_xml(sheet_names: Sequence[str]) -> str:
    sheets_xml = "".join(
        f'<sheet name="{xml_escape(name)}" sheetId="{idx}" r:id="rId{idx}"/>'
        for idx, name in enumerate(sheet_names, start=1)
    )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f"<sheets>{sheets_xml}</sheets>"
        "</workbook>"
    )


def _excel_workbook_rels(count: int) -> str:
    relations = "".join(
        f'<Relationship Id="rId{i}" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
        f'Target="worksheets/sheet{i}.xml"/>'
        for i in range(1, count + 1)
    )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        f"{relations}</Relationships>"
    )


def _excel_styles_xml() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        '<fonts count="1"><font><sz val="11"/><name val="Calibri"/></font></fonts>'
        '<fills count="1"><fill><patternFill patternType="none"/></fill></fills>'
        '<borders count="1"><border><left/><right/><top/><bottom/><diagonal/></border></borders>'
        '<cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>'
        '<cellXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/></cellXfs>'
        '</styleSheet>'
    )


def _excel_core_props() -> str:
    generated = datetime.now(timezone.utc).isoformat()
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
        'xmlns:dc="http://purl.org/dc/elements/1.1/" '
        'xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" '
        'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
        '<dc:title>Pharmaceuticals Financial Report</dc:title>'
        '<dc:creator>Pharmaceuticals Financial Model</dc:creator>'
        f'<dcterms:created xsi:type="dcterms:W3CDTF">{generated}</dcterms:created>'
        f'<dcterms:modified xsi:type="dcterms:W3CDTF">{generated}</dcterms:modified>'
        '</cp:coreProperties>'
    )


def _excel_app_props(sheet_names: Sequence[str]) -> str:
    parts = "".join(f"<vt:lpstr>{xml_escape(name)}</vt:lpstr>" for name in sheet_names)
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" '
        'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
        f'<HeadingPairs><vt:vector size="1" baseType="variant">'
        '<vt:variant><vt:lpstr>Worksheets</vt:lpstr></vt:variant>'
        '</vt:vector></HeadingPairs>'
        f'<TitlesOfParts><vt:vector size="{len(sheet_names)}" baseType="lpstr">{parts}</vt:vector></TitlesOfParts>'
        '</Properties>'
    )


def _build_word_fallback(blocks: Sequence[Tuple[str, str]]) -> bytes:
    buffer = BytesIO()
    with ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", _docx_content_types())
        archive.writestr("_rels/.rels", _docx_root_rels())
        archive.writestr("docProps/core.xml", _docx_core_props())
        archive.writestr("docProps/app.xml", _docx_app_props())
        archive.writestr("word/_rels/document.xml.rels", _docx_document_rels())
        archive.writestr("word/document.xml", _docx_document_xml(blocks))
    return buffer.getvalue()


def _docx_document_xml(blocks: Sequence[Tuple[str, str]]) -> str:
    parts = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">',
        "<w:body>",
    ]

    for kind, text in blocks:
        parts.append(_docx_paragraph_xml(kind, text))

    parts.append(
        '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" '
        'w:bottom="1440" w:left="1440"/></w:sectPr>'
    )
    parts.append("</w:body></w:document>")
    return "".join(parts)


def _docx_paragraph_xml(kind: str, text: str) -> str:
    style_map = {
        "title": "Title",
        "subtitle": "Subtitle",
        "heading1": "Heading1",
        "heading2": "Heading2",
        "note": "IntenseQuote",
    }
    style = style_map.get(kind)
    escaped = xml_escape(text)
    escaped = escaped.replace("\n", "&#10;")
    p_parts = ["<w:p>"]
    if style:
        p_parts.append(f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>')
    p_parts.append(f'<w:r><w:t xml:space="preserve">{escaped}</w:t></w:r></w:p>')
    return "".join(p_parts)


def _docx_content_types() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
        '<Override PartName="/docProps/app.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
        '</Types>'
    )


def _docx_root_rels() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/>'
        '<Relationship Id="rId2" '
        'Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" '
        'Target="docProps/core.xml"/>'
        '<Relationship Id="rId3" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" '
        'Target="docProps/app.xml"/>'
        '</Relationships>'
    )


def _docx_document_rels() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>'
    )


def _docx_core_props() -> str:
    generated = datetime.now(timezone.utc).isoformat()
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
        'xmlns:dc="http://purl.org/dc/elements/1.1/" '
        'xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" '
        'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
        '<dc:title>Pharmaceuticals Financial Report</dc:title>'
        '<dc:creator>Pharmaceuticals Financial Model</dc:creator>'
        f'<dcterms:created xsi:type="dcterms:W3CDTF">{generated}</dcterms:created>'
        f'<dcterms:modified xsi:type="dcterms:W3CDTF">{generated}</dcterms:modified>'
        '</cp:coreProperties>'
    )


def _docx_app_props() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" '
        'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
        '<Application>Pharmaceuticals Financial Model</Application>'
        '</Properties>'
    )


def _build_pdf_fallback(blocks: Sequence[Tuple[str, str]]) -> bytes:
    lines = _blocks_to_pdf_lines(blocks)
    if not lines:
        lines = ["Pharmaceuticals Financial Report"]

    line_height = 14
    page_height = 842
    margin = 36
    lines_per_page = max(int((page_height - 2 * margin) / line_height), 1)

    pages: List[List[str]] = []
    current: List[str] = []
    for line in lines:
        if len(current) >= lines_per_page:
            pages.append(current)
            current = []
        current.append(line)
    if current:
        pages.append(current)

    return _render_pdf_pages(pages, line_height=line_height, margin=margin, page_height=page_height)


def _blocks_to_pdf_lines(blocks: Sequence[Tuple[str, str]]) -> List[str]:
    lines: List[str] = []
    for kind, text in blocks:
        if kind == "title":
            lines.append(text)
            lines.append("")
        elif kind == "subtitle":
            lines.append(text)
            lines.append("")
        elif kind == "heading1":
            lines.append(text.upper())
        elif kind == "heading2":
            lines.append(f"  {text}")
        else:
            lines.append(text)
    return [line for line in lines if line is not None]


def _render_pdf_pages(
    pages: Sequence[Sequence[str]], *, line_height: int, margin: int, page_height: int
) -> bytes:
    page_count = max(len(pages), 1)
    catalog_id = 1
    pages_id = 2
    page_ids = [idx for idx in range(3, 3 + page_count)]
    content_ids = [idx for idx in range(3 + page_count, 3 + 2 * page_count)]
    font_id = 3 + 2 * page_count
    max_id = font_id

    objects: dict[int, bytes] = {}

    font_obj = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"
    objects[font_id] = font_obj

    page_refs = " ".join(f"{pid} 0 R" for pid in page_ids) or "3 0 R"
    pages_obj = f"<< /Type /Pages /Count {page_count} /Kids [{page_refs}] >>".encode("ascii")
    objects[pages_id] = pages_obj

    for idx, page in enumerate(pages or [["No data available."]]):
        page_id = page_ids[idx]
        content_id = content_ids[idx]
        stream = _pdf_content_stream(page, line_height=line_height, margin=margin, page_height=page_height)
        content = (
            f"<< /Length {len(stream)} >>\nstream\n".encode("ascii")
            + stream
            + b"\nendstream"
        )
        objects[content_id] = content
        page_obj = (
            f"<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 595 842] "
            f"/Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >>"
        ).encode("ascii")
        objects[page_id] = page_obj

    catalog_obj = f"<< /Type /Catalog /Pages {pages_id} 0 R >>".encode("ascii")
    objects[catalog_id] = catalog_obj

    buffer = BytesIO()
    buffer.write(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")

    offsets: dict[int, int] = {}
    for obj_id in range(1, max_id + 1):
        offsets[obj_id] = buffer.tell()
        buffer.write(f"{obj_id} 0 obj\n".encode("ascii"))
        buffer.write(objects.get(obj_id, b"<<>>"))
        buffer.write(b"\nendobj\n")

    startxref = buffer.tell()
    buffer.write(b"xref\n")
    buffer.write(f"0 {max_id + 1}\n".encode("ascii"))
    buffer.write(b"0000000000 65535 f \n")
    for obj_id in range(1, max_id + 1):
        offset = offsets[obj_id]
        buffer.write(f"{offset:010d} 00000 n \n".encode("ascii"))
    buffer.write(b"trailer\n")
    buffer.write(f"<< /Size {max_id + 1} /Root {catalog_id} 0 R >>\n".encode("ascii"))
    buffer.write(b"startxref\n")
    buffer.write(f"{startxref}\n".encode("ascii"))
    buffer.write(b"%%EOF")
    return buffer.getvalue()


def _pdf_content_stream(lines: Sequence[str], *, line_height: int, margin: int, page_height: int) -> bytes:
    y = page_height - margin
    commands = ["BT", "/F1 12 Tf"]
    for line in lines:
        escaped = _pdf_escape(line)
        commands.append(f"1 0 0 1 {margin} {int(y)} Tm ({escaped}) Tj")
        y -= line_height
    commands.append("ET")
    return "\n".join(commands).encode("latin-1")


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _ordered_columns(rows: Sequence[Mapping[str, Any]]) -> List[str]:
    seen: List[str] = []
    for row in rows:
        for key in row.keys():
            if key not in seen:
                seen.append(key)
    return seen

def _table_to_rows(data: Any) -> List[Mapping[str, Any]]:
    if data is None:
        return []
    if isinstance(data, list):
        if data and isinstance(data[0], Mapping):
            return [dict(row) for row in data]  # shallow copy
        return [
            {"Value": item}
            for item in data
        ]
    if isinstance(data, Table):
        mapping = data.as_dict()
        rows: List[MutableMapping[str, Any]] = []
        for position, idx in enumerate(data.index):
            row: MutableMapping[str, Any] = {data.index_name: idx}
            for column, values in mapping.items():
                row[column] = values[position]
            rows.append(row)
        return rows
    if pd is not None and isinstance(data, pd.DataFrame):
        frame = data.reset_index()
        return frame.to_dict(orient="records")
    if isinstance(data, Mapping):
        return [
            {"Metric": key, "Value": value}
            for key, value in data.items()
        ]
    try:
        return json.loads(json.dumps(data))
    except Exception:
        return [{"Value": _stringify(data)}]


def _table_to_dataframe(data: Any):
    if pd is None:
        return None
    if isinstance(data, pd.DataFrame):
        return data
    if isinstance(data, Table):
        return data.to_frame()
    rows = _table_to_rows(data)
    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


def _escape_csv(value: Any) -> str:
    text = _stringify(value)
    if any(char in text for char in [",", "\n", '"']):
        text = text.replace('"', '""')
        return f'"{text}"'
    return text


def _sheet_name(section: str, table: str, tracker: MutableMapping[str, int]) -> str:
    base = f"{section} {table}".strip() or "Sheet"
    base = re.sub(r"[^A-Za-z0-9 ]", "", base)
    base = re.sub(r"\s+", " ", base).strip()
    if not base:
        base = "Sheet"
    base = base[:31]
    count = tracker.get(base, 0)
    if count:
        suffix = f"_{count}"
        base = f"{base[:31-len(suffix)]}{suffix}"
    tracker[base] = count + 1
    return base


def _stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        if math.isnan(value):
            return ""
        return f"{value:,.4f}" if abs(value) < 1 else f"{value:,.2f}"
    return str(value)


def _pdf_multiline(pdf: "FPDF", text: str, *, bold: bool = False) -> None:
    if bold:
        pdf.set_font("Arial", "B", 10)
    else:
        pdf.set_font("Arial", "", 10)
    pdf.multi_cell(0, 6, text)
    if bold:
        pdf.set_font("Arial", "", 10)


__all__ = [
    "REPORT_FORMATS",
    "ReportGenerationError",
    "ReportSection",
    "ReportTable",
    "collect_report_sections",
    "generate_report",
]
